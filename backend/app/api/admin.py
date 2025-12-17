"""
FinanceClinics - Admin Dashboard API
"""

from datetime import datetime, timedelta
from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from sqlalchemy import func
from ..models import User, Page, Service, BlogPost, Lead
from ..models import MISTemplate, MISData
from ..extensions import db

admin_bp = Blueprint('admin', __name__)


@admin_bp.route('/dashboard', methods=['GET'])
@jwt_required()
def dashboard():
    """Get dashboard statistics"""
    # Total counts
    total_leads = Lead.query.count()
    total_pages = Page.query.count()
    total_services = Service.query.count()
    total_posts = BlogPost.query.count()
    
    # Recent leads (last 7 days)
    week_ago = datetime.utcnow() - timedelta(days=7)
    recent_leads = Lead.query.filter(Lead.created_at >= week_ago).count()
    
    # New leads (unread)
    new_leads = Lead.query.filter_by(status='new').count()
    
    # Leads by status
    leads_by_status = db.session.query(
        Lead.status, func.count(Lead.id)
    ).group_by(Lead.status).all()
    
    # Recent blog views
    total_views = db.session.query(func.sum(BlogPost.views)).scalar() or 0
    
    # Monthly leads (last 6 months)
    six_months_ago = datetime.utcnow() - timedelta(days=180)
    monthly_leads = db.session.query(
        func.date_format(Lead.created_at, '%Y-%m').label('month'),
        func.count(Lead.id)
    ).filter(Lead.created_at >= six_months_ago)\
     .group_by('month')\
     .order_by('month')\
     .all()
    
    # Recent activity
    recent_activity = []
    
    # Recent leads
    latest_leads = Lead.query.order_by(Lead.created_at.desc()).limit(5).all()
    for lead in latest_leads:
        recent_activity.append({
            'type': 'lead',
            'message': f'New inquiry from {lead.name}',
            'time': lead.created_at.isoformat(),
            'id': lead.id
        })
    
    # Recent posts
    latest_posts = BlogPost.query.order_by(BlogPost.created_at.desc()).limit(3).all()
    for post in latest_posts:
        recent_activity.append({
            'type': 'post',
            'message': f'Blog post: {post.title}',
            'time': post.created_at.isoformat(),
            'id': post.id
        })
    
    # Sort by time
    recent_activity.sort(key=lambda x: x['time'], reverse=True)
    
    return jsonify({
        'stats': {
            'total_leads': total_leads,
            'total_pages': total_pages,
            'total_services': total_services,
            'total_posts': total_posts,
            'recent_leads': recent_leads,
            'new_leads': new_leads,
            'total_views': total_views
        },
        'leads_by_status': {status: count for status, count in leads_by_status},
        'monthly_leads': [{'month': month, 'count': count} for month, count in monthly_leads],
        'recent_activity': recent_activity[:10]
    }), 200


@admin_bp.route('/users', methods=['GET'])
@jwt_required()
def get_users():
    """Get all admin users"""
    users = User.query.all()
    return jsonify({
        'users': [u.to_dict() for u in users]
    }), 200


@admin_bp.route('/users', methods=['POST'])
@jwt_required()
def create_user():
    """Create new admin user"""
    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    name = data.get('name', '').strip()
    
    if not email or not password or not name:
        return jsonify({'error': 'Email, password, and name are required'}), 400
    
    if len(password) < 8:
        return jsonify({'error': 'Password must be at least 8 characters'}), 400
    
    if User.query.filter_by(email=email).first():
        return jsonify({'error': 'User with this email already exists'}), 400
    
    user = User(
        email=email,
        name=name,
        role=data.get('role', 'admin')
    )
    user.set_password(password)
    
    db.session.add(user)
    db.session.commit()
    
    return jsonify({'message': 'User created', 'user': user.to_dict()}), 201


@admin_bp.route('/users/<int:user_id>', methods=['PUT'])
@jwt_required()
def update_user(user_id):
    """Update admin user"""
    user = User.query.get_or_404(user_id)
    data = request.get_json()
    
    if 'name' in data:
        user.name = data['name']
    if 'email' in data:
        new_email = data['email'].strip().lower()
        if new_email != user.email:
            if User.query.filter_by(email=new_email).first():
                return jsonify({'error': 'Email already in use'}), 400
            user.email = new_email
    if 'is_active' in data:
        user.is_active = data['is_active']
    if 'password' in data and data['password']:
        if len(data['password']) < 8:
            return jsonify({'error': 'Password must be at least 8 characters'}), 400
        user.set_password(data['password'])
    
    db.session.commit()
    
    return jsonify({'message': 'User updated', 'user': user.to_dict()}), 200


@admin_bp.route('/users/<int:user_id>', methods=['DELETE'])
@jwt_required()
def delete_user(user_id):
    """Delete admin user"""
    current_user_id = get_jwt_identity()
    
    if user_id == current_user_id:
        return jsonify({'error': 'Cannot delete your own account'}), 400
    
    user = User.query.get_or_404(user_id)
    
    db.session.delete(user)
    db.session.commit()
    
    return jsonify({'message': 'User deleted'}), 200


@admin_bp.route('/users/<int:user_id>/approve', methods=['PUT'])
@jwt_required()
def approve_user(user_id):
    """Approve (activate) a user account"""
    user = User.query.get_or_404(user_id)
    user.is_active = True
    db.session.commit()
    return jsonify({'message': 'User approved', 'user': user.to_dict()}), 200


# --- MIS Templates and Data ---


@admin_bp.route('/mis/templates', methods=['GET'])
@jwt_required()
def list_mis_templates():
    try:
        templates = MISTemplate.query.order_by(MISTemplate.created_at.desc()).all()
        return jsonify({'templates': [t.to_dict() for t in templates]}), 200
    except Exception:
        # If table doesn't exist or any DB error happens, return empty list
        return jsonify({'templates': []}), 200


@admin_bp.route('/mis/templates', methods=['POST'])
@jwt_required()
def create_mis_template():
    data = request.get_json() or {}
    name = data.get('name')
    columns = data.get('columns', [])
    if not name:
        return jsonify({'error': 'Template name required'}), 400
    import json
    tpl = MISTemplate(name=name, columns=json.dumps(columns), created_by=get_jwt_identity())
    db.session.add(tpl)
    db.session.commit()
    return jsonify({'template': tpl.to_dict()}), 201


@admin_bp.route('/mis/templates/<int:tpl_id>', methods=['GET'])
@jwt_required()
def get_mis_template(tpl_id):
    tpl = MISTemplate.query.get_or_404(tpl_id)
    return jsonify({'template': tpl.to_dict()}), 200


@admin_bp.route('/mis/templates/<int:tpl_id>', methods=['PUT'])
@jwt_required()
def update_mis_template(tpl_id):
    tpl = MISTemplate.query.get_or_404(tpl_id)
    data = request.get_json() or {}
    name = data.get('name')
    columns = data.get('columns')
    import json
    if name:
        tpl.name = name
    if columns is not None:
        tpl.columns = json.dumps(columns)
    db.session.commit()
    return jsonify({'template': tpl.to_dict()}), 200


@admin_bp.route('/mis/templates/<int:tpl_id>', methods=['DELETE'])
@jwt_required()
def delete_mis_template(tpl_id):
    tpl = MISTemplate.query.get_or_404(tpl_id)
    # delete associated data rows
    MISData.query.filter_by(template_id=tpl.id).delete()
    db.session.delete(tpl)
    db.session.commit()
    return jsonify({'message': 'Template and data deleted'}), 200


@admin_bp.route('/mis/templates/<int:tpl_id>/rows', methods=['GET'])
@jwt_required()
def list_mis_rows(tpl_id):
    tpl = MISTemplate.query.get_or_404(tpl_id)
    rows = MISData.query.filter_by(template_id=tpl.id).order_by(MISData.id.desc()).all()
    return jsonify({'rows': [r.to_dict() for r in rows]}), 200


@admin_bp.route('/mis/templates/<int:tpl_id>/import', methods=['POST'])
@jwt_required()
def import_mis_data(tpl_id):
    tpl = MISTemplate.query.get_or_404(tpl_id)
    # Only accept file uploads
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    file = request.files['file']
    filename = file.filename or ''
    # support CSV import
    if filename.lower().endswith('.csv') or request.form.get('format') == 'csv':
        import csv, io, json
        stream = io.StringIO(file.stream.read().decode('utf-8', errors='replace'))
        reader = csv.DictReader(stream)
        count = 0
        for row in reader:
            # keep keys that match template columns if possible
            try:
                data_obj = {k: v for k, v in row.items()}
                mis_row = MISData(template_id=tpl.id, data=json.dumps(data_obj))
                db.session.add(mis_row)
                count += 1
            except Exception:
                continue
        db.session.commit()
        return jsonify({'imported': count}), 200
    # Attempt DOCX import if python-docx available
    if filename.lower().endswith('.docx') or request.form.get('format') == 'docx':
        try:
            from docx import Document
        except Exception:
            Document = None
        if Document:
            import json
            doc = Document(file)
            count = 0
            # iterate tables and import rows
            for table in doc.tables:
                # assume first row is header
                headers = [c.text.strip() for c in table.rows[0].cells]
                for row in table.rows[1:]:
                    vals = [c.text for c in row.cells]
                    data_obj = {h: v for h, v in zip(headers, vals)}
                    try:
                        mis_row = MISData(template_id=tpl.id, data=json.dumps(data_obj))
                        db.session.add(mis_row)
                        count += 1
                    except Exception:
                        continue
            db.session.commit()
            return jsonify({'imported': count}), 200

    # Attempt Excel if pandas available
    try:
        import pandas as pd
    except Exception:
        pd = None
    if pd and (filename.lower().endswith(('.xls', '.xlsx')) or request.form.get('format') in ('xls', 'xlsx', 'excel')):
        import json
        df = pd.read_excel(file)
        count = 0
        for _, row in df.iterrows():
            data_obj = row.where(pd.notnull(row), None).to_dict()
            mis_row = MISData(template_id=tpl.id, data=json.dumps(data_obj))
            db.session.add(mis_row)
            count += 1
        db.session.commit()
        return jsonify({'imported': count}), 200

    # Attempt PDF import using pdfplumber (extract table data)
    if filename.lower().endswith('.pdf') or request.form.get('format') == 'pdf':
        try:
            import pdfplumber
        except Exception:
            pdfplumber = None
        if pdfplumber:
            import io, json
            stream = io.BytesIO(file.stream.read())
            count = 0
            try:
                with pdfplumber.open(stream) as pdf:
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        for table in tables:
                            if not table or len(table) < 2:
                                continue
                            headers = [c.strip() if c else f'col{i}' for i, c in enumerate(table[0])]
                            for row in table[1:]:
                                data_obj = {headers[i]: (row[i] or '') for i in range(min(len(headers), len(row)))}
                                try:
                                    mis_row = MISData(template_id=tpl.id, data=json.dumps(data_obj))
                                    db.session.add(mis_row)
                                    count += 1
                                except Exception:
                                    continue
                db.session.commit()
                return jsonify({'imported': count}), 200
            except Exception:
                return jsonify({'error': 'Failed to parse PDF'}), 500

    return jsonify({'error': 'Unsupported file format or missing dependency for Excel import'}), 501


@admin_bp.route('/mis/templates/<int:tpl_id>/export', methods=['GET'])
@jwt_required()
def export_mis_data(tpl_id):
    tpl = MISTemplate.query.get_or_404(tpl_id)
    fmt = request.args.get('format', 'csv').lower()
    rows = MISData.query.filter_by(template_id=tpl.id).all()
    import json, io, csv
    if fmt == 'csv':
        output = io.StringIO()
        # determine headers from template columns keys if available
        try:
            cols = json.loads(tpl.columns or '[]')
            headers = [c.get('key') for c in cols if c.get('key')]
        except Exception:
            headers = []
        writer = csv.writer(output)
        if headers:
            writer.writerow(headers)
            for r in rows:
                obj = json.loads(r.data or '{}')
                writer.writerow([obj.get(h, '') for h in headers])
        else:
            # fallback: export each row as JSON string per line
            writer.writerow(['data'])
            for r in rows:
                writer.writerow([r.data])
        output.seek(0)
        return (output.getvalue(), 200, {
            'Content-Type': 'text/csv',
            'Content-Disposition': f'attachment; filename="{tpl.name}.csv"'
        })

    # Excel via pandas if available
    try:
        import pandas as pd
    except Exception:
        pd = None
    if pd and fmt in ('xls', 'xlsx', 'excel'):
        all_objs = []
        for r in rows:
            try:
                all_objs.append(json.loads(r.data or '{}'))
            except Exception:
                all_objs.append({})
        df = pd.DataFrame(all_objs)
        output = io.BytesIO()
        # Use openpyxl engine if available
        try:
            df.to_excel(output, index=False, engine='openpyxl')
        except Exception:
            df.to_excel(output, index=False)
        output.seek(0)
        return send_file(
            output,
            as_attachment=True,
            download_name=f"{tpl.name}.xlsx",
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    # DOCX export using python-docx
    if fmt == 'docx':
        try:
            from docx import Document
            import json, io
        except Exception:
            Document = None
        if Document:
            doc = Document()
            try:
                cols = json.loads(tpl.columns or '[]')
                headers = [c.get('key') for c in cols if c.get('key')]
            except Exception:
                headers = []
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = h
                for r in rows:
                    try:
                        obj = json.loads(r.data or '{}')
                    except Exception:
                        obj = {}
                    row_cells = table.add_row().cells
                    for i, h in enumerate(headers):
                        row_cells[i].text = str(obj.get(h, '') or '')
            else:
                doc.add_paragraph('No structured headers; exported rows follow as JSON:')
                for r in rows:
                    doc.add_paragraph(r.data)
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            return send_file(
                bio,
                as_attachment=True,
                download_name=f"{tpl.name}.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

    # PDF export: simple table via reportlab
    if fmt == 'pdf':
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.platypus import Table, SimpleDocTemplate
            from reportlab.lib import colors
            import json, io
        except Exception:
            SimpleDocTemplate = None
        if SimpleDocTemplate:
            all_objs = []
            try:
                cols = json.loads(tpl.columns or '[]')
                headers = [c.get('key') for c in cols if c.get('key')]
            except Exception:
                headers = []
            if headers:
                table_data = [headers]
                for r in rows:
                    try:
                        obj = json.loads(r.data or '{}')
                    except Exception:
                        obj = {}
                    table_data.append([str(obj.get(h, '') or '') for h in headers])
            else:
                table_data = [['data']]
                for r in rows:
                    table_data.append([r.data])

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=letter)
            tbl = Table(table_data, repeatRows=1)
            tbl.setStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.gray),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ])
            doc.build([tbl])
            buf.seek(0)
            return send_file(
                buf,
                as_attachment=True,
                download_name=f"{tpl.name}.pdf",
                mimetype='application/pdf'
            )

    return jsonify({'error': 'Export format not supported'}), 501
